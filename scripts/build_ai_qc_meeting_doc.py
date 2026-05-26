from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"E:\yolo\AI_QC_meeting_checklist.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 30, 45)
MUTED = RGBColor(90, 99, 112)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF7D6"
PALE_GREEN = "EAF6EA"
PALE_RED = "FCECEC"
WHITE = "FFFFFF"
BORDER = "C8D2E0"


def set_font(run, size=None, bold=None, color=None, name="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=10.5, color=INK, bold=False):
    for run in paragraph.runs:
        set_font(run, size=size, bold=bold, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def set_borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_text(cell, text, size=9.3, bold=False, color=INK, fill=None, align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    if fill:
        set_cell_shading(cell, fill)


def add_para(doc, text="", size=10.5, bold=False, color=INK, after=6, before=0, align=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.style = f"Heading {level}"
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color=BLUE)
        paragraph.paragraph_format.space_before = Pt(18)
        paragraph.paragraph_format.space_after = Pt(10)
    elif level == 2:
        set_font(run, size=13, bold=True, color=BLUE)
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after = Pt(7)
    else:
        set_font(run, size=12, bold=True, color=DARK_BLUE)
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after = Pt(5)
    return paragraph


def add_callout(doc, label, body, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_borders(table, color="E0C45F", size="6")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.2
    r1 = paragraph.add_run(label + "：")
    set_font(r1, size=10.5, bold=True, color=DARK_BLUE)
    r2 = paragraph.add_run(body)
    set_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_check_table(doc, title, rows):
    add_heading(doc, title, 2)
    widths = [720, 2280, 3480, 2880]
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, widths)
    set_borders(table)
    headers = ["确认", "事项", "会上要问", "记录/结论"]
    for idx, header in enumerate(headers):
        cell_text(table.cell(0, idx), header, size=9.5, bold=True, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item, question, note in rows:
        cells = table.add_row().cells
        cell_text(cells[0], "□", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(cells[1], item, size=9.1, bold=True)
        cell_text(cells[2], question, size=9.1)
        cell_text(cells[3], note, size=9.1, color=MUTED)
    set_table_geometry(table, widths)
    set_borders(table)


def add_matrix(doc, title, headers, rows, widths):
    add_heading(doc, title, 2)
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_borders(table)
    for idx, header in enumerate(headers):
        cell_text(table.cell(0, idx), header, size=9.2, bold=True, fill=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value, size=8.9)
    set_table_geometry(table, widths)
    set_borders(table)
    return table


def add_scenario(doc, name, position, shooting, questions, first_delivery):
    add_heading(doc, name, 2)
    table = doc.add_table(rows=4, cols=2)
    widths = [1980, 7380]
    set_table_geometry(table, widths)
    set_borders(table)
    data = [
        ("位置关系", position),
        ("拍摄引导", shooting),
        ("会上必须问清", questions),
        ("首版交付边界", first_delivery),
    ]
    for idx, (label, value) in enumerate(data):
        cell_text(table.cell(idx, 0), label, size=9.3, bold=True, fill=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(table.cell(idx, 1), value, size=9.1)
    set_table_geometry(table, widths)
    set_borders(table)


def add_signature_table(doc):
    add_heading(doc, "七、会后动作与责任分工", 1)
    rows = [
        ["确定首批试点场景", "建设公司/我方共同确认", "会后当天", "□ 已确认 □ 待补充"],
        ["提供现场样本", "建设公司", "1-2个工作日", "每个场景30-50张照片或5-10段短视频"],
        ["整理拍摄规范", "我方", "1个工作日", "形成现场采集说明"],
        ["制作首版Demo", "我方", "按数据到位时间评估", "先离线识别，再评估平台接入"],
        ["确认验收口径", "建设公司/我方", "Demo前", "误报、漏报、人工复核流程"],
    ]
    add_matrix(
        doc,
        "行动项清单",
        ["行动项", "责任方", "时间", "备注"],
        rows,
        [2520, 2160, 1800, 2880],
    )
    add_heading(doc, "会议记录区", 2)
    for label in ["最终确定首批场景", "现场数据联系人", "需要建设公司补充的规范/图纸/样本", "我方会后输出物", "其他风险或限制"]:
        p = add_para(doc, f"{label}：", size=10.2, bold=True, color=DARK_BLUE, after=2)
        for _ in range(2):
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(2)
            run = line.add_run("____________________________________________________________________________")
            set_font(run, size=9, color=RGBColor(170, 176, 184))


def setup_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = "AI质监需求对齐会议确认清单"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_font(header, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "内部会议使用 | 2026年5月20日"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_font(footer, size=9, color=MUTED)


def build():
    doc = Document()
    setup_doc(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("AI质监需求对齐会议确认清单")
    set_font(run, size=24, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("建设公司沟通版 | 重点明确首批场景、拍摄规范、数据要求与首版交付边界")
    set_font(run, size=12, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [2160, 7200])
    set_borders(meta)
    for idx, (label, value) in enumerate([
        ("会议日期", "2026年5月20日"),
        ("会议目标", "从质量管理AI需求中选出2-4个可快速验证的场景，并明确现场采集、识别规则、定位闭环和交付标准。"),
        ("建议主线", "先做手机拍照/短视频可完成、规则清楚、样本容易获取的场景；复杂测量和BIM一致性类需求放到后续批次。"),
        ("预期产出", "首批场景清单、每个场景的识别范围、拍摄规范、样本清单、Demo/POC交付形式、责任人和时间。"),
    ]):
        cell_text(meta.cell(idx, 0), label, size=9.5, bold=True, fill=LIGHT_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(meta.cell(idx, 1), value, size=9.3)
    add_callout(
        doc,
        "开场可直接这样说",
        "本次会议先不讨论所有AI质检需求全部落地，重点是确定首批2-4个高频、规则清楚、拍照可实现的试点场景，并把拍摄方式、判定规则、数据样本和首版交付标准确认下来。",
        PALE_GREEN,
    )

    add_heading(doc, "一、会议要明确的最终结论", 1)
    add_check_table(doc, "结论清单", [
        ("首批试点场景", "最终选哪2-4个？是否按“砼缺陷、支架顶托/底托、钢筋间距、材料计数”推进？", "记录最终场景名称和优先级。"),
        ("每个场景识别对象", "具体识别哪些问题？哪些先不做？例如裂缝宽度、石材平整度、焊缝质量是否放后续。", "避免需求泛化成“什么都识别”。"),
        ("合格/不合格规则", "判定依据来自规范、企业标准还是项目经验？阈值由谁提供？", "没有规则就先做疑似问题提示。"),
        ("拍照/视频采集方式", "谁拍、用什么设备、距离角度、是否全景+近景、是否带尺/标定板？", "这是后续效果的关键。"),
        ("现场定位字段", "结果是否必须关联项目、楼栋、楼层、轴线、构件编号、照片点位？", "决定是否能进入整改闭环。"),
        ("样本和标注", "是否有历史照片/视频？每类场景能否提供正负样本？谁负责确认标注？", "确定Demo能否快速启动。"),
        ("首版交付形式", "要离线Demo、手机拍照识别、问题清单，还是接入既有平台？", "先定最小可交付版本。"),
        ("验收口径", "误报、漏报、人工复核如何接受？Demo阶段看什么指标？", "避免后期评价标准不一致。"),
    ])

    add_heading(doc, "二、建议优先讨论的场景", 1)
    add_matrix(
        doc,
        "首批场景建议",
        ["优先级", "场景", "为什么适合先做", "会议建议结论"],
        [
            ["1", "砼结构表面缺陷自动查找", "视觉特征直观，照片样本容易采集，能快速展示AI框选缺陷和问题清单。", "建议作为首批必选，先识别蜂窝、麻面、裂缝、露筋等明显问题。"],
            ["2", "支架顶托/底托质量检测", "规则相对明确，位置关系清楚，适合用手机照片做快速核查。", "建议作为首批必选，先做外露长度、螺母状态、接触关系等可见问题。"],
            ["3", "钢筋工程数量、间距检测", "质量管控高频，结果可视化强，但需要尺度标定和拍摄规范配合。", "建议列为首批或第二优先，先限定平面钢筋网/规则区域。"],
            ["4", "材料进场数量识别", "目标检测计数成熟，适合演示自动盘点，但要明确计根、计捆还是计批次。", "可作为备选，时间足够时加入。"],
        ],
        [960, 2160, 3960, 2280],
    )
    add_callout(
        doc,
        "暂缓建议",
        "盾构收敛变形、石材平整度/高低差、焊缝质量、BIM符合性识别等价值高，但对尺度标定、传感器、拍摄标准和数据样本要求更高，建议放到后续批次。",
        PALE_YELLOW,
    )

    add_heading(doc, "三、所有场景都要问清的共性问题", 1)
    add_check_table(doc, "共性问题清单", [
        ("业务边界", "这个场景人工质检现在怎么做？检查频率、责任人和整改流程是什么？", "决定AI插入哪个环节。"),
        ("对象范围", "识别对象是否限定在某类构件、某个施工阶段、某种拍摄视角？", "范围越清楚，越容易快速交付。"),
        ("判定标准", "哪些情况一定报警？哪些只是提示人工复核？阈值是否有书面依据？", "先区分“识别”和“判定”。"),
        ("位置关系", "异常和构件、轴线、支撑件、相邻钢筋之间的关系怎么定义？", "涉及测距、间距和构件定位。"),
        ("拍摄规范", "需要全景、近景还是视频扫拍？是否必须正拍？是否允许斜拍？", "拍摄不统一会直接影响识别。"),
        ("尺度来源", "需要算长度、面积、间距时，尺度来自钢尺、标定板、已知构件尺寸还是AR测距？", "没有尺度就不要承诺精确尺寸。"),
        ("数据闭环", "AI结果输出给谁？是问题清单、整改通知、统计报表，还是接入现有系统？", "决定交付界面和字段。"),
        ("安全合规", "现场视频照片是否可用于模型训练？是否需要脱敏？是否允许外发？", "提前规避数据使用风险。"),
    ])

    add_heading(doc, "四、首批场景逐项确认", 1)
    add_scenario(
        doc,
        "1. 砼结构表面缺陷自动查找",
        "要明确缺陷相对构件的位置，例如墙、柱、梁、板的构件编号、楼层、轴线、边角/中部位置；如果要估算缺陷面积，需要画面中有明确尺度参照。",
        "建议全景+近景两张：全景用于定位构件，近景用于识别缺陷。近景尽量正对表面，避免强反光、阴影、水渍遮挡；如需面积估算，照片内放钢尺、标定板或其他已知尺寸参照物。",
        "优先识别哪些缺陷：蜂窝、麻面、裂缝、露筋、孔洞、掉角？裂缝宽度/长度、蜂窝面积是否有报警阈值？首版是否接受只输出“疑似缺陷+人工复核”？",
        "首版建议做离线图片识别Demo：框选缺陷位置、输出缺陷类型、置信度和问题截图。面积和裂缝宽度先作为可选能力，依赖标尺/标定板。",
    )
    add_scenario(
        doc,
        "2. 支架顶托/底托质量检测",
        "核心是立杆、顶托/底托、丝杆、螺母、主楞/垫板之间的关系。要判断丝杆外露长度、螺母旋合状态、顶托是否有效支撑、底托是否落在垫板或基础上。",
        "建议侧向或45度拍摄，画面必须包含立杆、顶托/底托、丝杆、螺母和接触面；如果要量外露长度，必须放入钢尺或使用固定距离标定。避免只拍局部螺母导致无法判断支撑关系。",
        "外露长度合格阈值是多少？不同型号顶托/底托是否标准不同？顶托偏斜、悬空、未顶紧、底托无垫板是否都算问题？哪些问题首版必须做，哪些后续做？",
        "首版建议先识别构件和明显异常：顶托/底托存在性、丝杆外露疑似超限、螺母状态异常、接触关系异常。精确测量在拍摄标定后再承诺。",
    )
    add_scenario(
        doc,
        "3. 钢筋工程数量、间距检测",
        "要明确相邻钢筋中心距、横纵向钢筋方向、检测区域边界，以及是否只针对规则平面钢筋网。复杂梁柱节点、密集搭接区、遮挡严重区域建议先排除。",
        "尽量正对钢筋平面拍摄，画面包含完整检测区域和边界；需要钢尺、标定板、已知模板尺寸或已知钢筋直径作为尺度参照。视频扫拍要保持速度稳定，并避免大角度透视。",
        "检测的是根数、间距、漏筋、保护层还是绑扎状态？允许偏差是多少？是否需要和图纸比对？首版能否先做单张照片内的数量和粗略间距检测？",
        "首版建议限定平面钢筋网：识别钢筋方向、根数、疑似间距超限位置。图纸比对、保护层厚度、复杂节点识别放到第二阶段。",
    )
    add_scenario(
        doc,
        "4. 材料进场数量识别",
        "要明确计数单位：根、捆、件、平方米、立方米还是批次。材料应尽量规则堆放，并与材料标牌、规格型号、进场单据建立关系。",
        "拍摄时需要覆盖完整堆放区域，尽量从正面或斜上方拍，减少遮挡；一张全景保留批次/标牌信息，一张近景展示材料端部或规格。堆放凌乱时只能做估算或提示人工复核。",
        "优先材料是哪几类：钢筋、钢管、方木、模板、管材？是否要求和进场单据自动比对？误差范围能接受多少？",
        "首版建议做可见目标计数：按照片/视频统计钢筋、管材或方木数量，输出数量、截图和人工复核提示。重量、体积和单据比对后续再做。",
    )

    add_heading(doc, "五、拍照/视频采集规范要点", 1)
    add_matrix(
        doc,
        "采集规范字段",
        ["类别", "必须明确", "建议要求"],
        [
            ["定位字段", "项目、楼栋、楼层、轴线、构件编号、施工区域、拍摄人、拍摄时间。", "现场二维码或表单填写，避免照片脱离位置后无法整改。"],
            ["画面类型", "全景用于定位，近景用于识别，复检照片用于闭环。", "重要问题尽量保留全景+近景两张。"],
            ["角度距离", "是否正拍、侧拍、斜拍；大致距离范围；是否允许视频扫拍。", "测量类场景尽量正拍，减少透视误差。"],
            ["尺度参照", "钢尺、标定板、已知构件尺寸、已知材料规格、AR测距或固定摄像头标定。", "没有尺度参照时只输出疑似识别，不输出精确尺寸。"],
            ["画质要求", "清晰、无遮挡、光照均匀，目标占画面比例足够。", "模糊、逆光、严重遮挡、过远照片应退回重拍。"],
            ["样本结构", "每个场景要有合格样本、问题样本、不同光照/距离/角度样本。", "首批每个场景先要30-50张照片或5-10段短视频。"],
        ],
        [1440, 4320, 3600],
    )
    add_callout(
        doc,
        "对外口径",
        "拍照规范不是额外负担，而是AI质检能否稳定落地的前提。第一版可以先把拍摄要求做轻：全景定位、近景识别、必要时带尺。",
        PALE_GREEN,
    )

    add_heading(doc, "六、首版交付边界建议", 1)
    add_matrix(
        doc,
        "MVP边界",
        ["事项", "首版建议做", "首版暂不承诺"],
        [
            ["识别方式", "离线Demo或手机上传图片/短视频后识别。", "现场实时全覆盖监控。"],
            ["输出内容", "问题截图、框选位置、问题类型、置信度、复核状态。", "所有尺寸都精确测量。"],
            ["规则判定", "有明确阈值的做合格/不合格；无阈值的做疑似问题提示。", "替代人工最终验收结论。"],
            ["数据闭环", "导出问题清单，可按项目/楼层/构件筛选。", "复杂流程系统深度集成。"],
            ["验收指标", "看识别效果、拍摄规范可执行性、误报/漏报是否可接受。", "直接用生产级准确率考核第一版Demo。"],
        ],
        [2160, 3600, 3600],
    )

    add_signature_table(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
