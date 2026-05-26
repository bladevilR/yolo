from __future__ import annotations

import html
import json
import math
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


WORKSPACE = Path("E:/yolo")
DELIVERY_DIR = WORKSPACE / "AI_QC_Human_Delivery_20260525"
CLEAR_ASSETS = DELIVERY_DIR / "clear_review_assets"

MANIFEST_PATH = WORKSPACE / "datasets/field_qc/reports/media_manifest.json"
MATERIAL_REPORT = WORKSPACE / "datasets/field_qc/rebar_material_counting/reports/rebar_material_count_report.json"
COUPLER_REPORT = WORKSPACE / "datasets/field_qc/rebar_coupler_thread_qc/reports/rebar_coupler_thread_qc_report.json"
CONCRETE_REPORT = WORKSPACE / "datasets/field_qc/concrete_surface_qc/reports/concrete_surface_qc_report.json"

HTML_OUT = DELIVERY_DIR / "04_AI_QC_Clear_Review_CN.html"
DOCX_OUT = DELIVERY_DIR / "04_AI_QC_Clear_Review_CN.docx"

COLORS = {
    "red": (218, 53, 42),
    "yellow": (232, 154, 33),
    "green": (42, 134, 90),
    "blue": (31, 96, 160),
    "dark": (25, 34, 44),
    "light_bg": (248, 250, 252),
    "white": (255, 255, 255),
}

CLASS_CN = {
    "repair_patch": "修补痕迹/表面补修疑似",
    "color_difference": "色差/污染疑似",
    "edge_defect": "边角缺损疑似",
    "surface_pitting": "麻面/起砂疑似",
    "crack_like": "裂缝疑似",
    "honeycomb_like": "蜂窝疑似",
}

COUPLER_REVIEWS = {
    "field-qc-0019": {
        "level": "yellow",
        "verdict": "不可作为完整接头验收",
        "title": "近景清楚，但更像单端/未完整连接状态",
        "bullets": [
            "这张不是清点套筒数量，重点是看连接是否完整。",
            "右侧未看到另一根钢筋进入套筒，像未完成连接或单端样件。",
            "结论：可用于识别套筒位置和拍照样例，不能作为合格接头判断。",
        ],
        "boxes": [(0.10, 0.48, 0.76, 0.78)],
    },
    "field-qc-0020": {
        "level": "yellow",
        "verdict": "只能定位，不能验收",
        "title": "远景多接头，套筒太小",
        "bullets": [
            "画面里有多个接头，但单个接头细节太小。",
            "能做“套筒定位/拍照引导”，不适合直接判露丝是否合规。",
            "结论：需要逐个接头近景补拍。",
        ],
        "boxes": [(0.39, 0.04, 0.58, 0.13), (0.38, 0.19, 0.60, 0.28), (0.36, 0.44, 0.61, 0.54), (0.37, 0.62, 0.63, 0.72)],
    },
    "field-qc-0021": {
        "level": "red",
        "verdict": "疑似露丝偏多，建议优先复核",
        "title": "近景可看，局部露丝较明显",
        "bullets": [
            "中间接头右侧外露螺纹比较明显，肉眼看有超出常规2扣/2p口径的风险。",
            "前景钢筋遮挡一部分边界，正式验收仍需正对、无遮挡近拍。",
            "结论：可作为“异常预警样本”进入第一版交付。",
        ],
        "boxes": [(0.14, 0.29, 0.88, 0.50), (0.26, 0.72, 0.82, 0.91)],
    },
    "field-qc-0022": {
        "level": "yellow",
        "verdict": "画面偏远，不能直接判",
        "title": "多接头中远景，细节不足",
        "bullets": [
            "能看到若干套筒，但接头边界和外露螺纹细节不够清楚。",
            "这类图适合做“自动发现接头位置”，不适合输出合格/不合格。",
            "结论：需要补拍单个接头近景。",
        ],
        "boxes": [(0.34, 0.27, 0.73, 0.74)],
    },
    "field-qc-0023": {
        "level": "red",
        "verdict": "疑似露丝偏多，建议优先复核",
        "title": "近景较清楚，两侧露丝可见",
        "bullets": [
            "中间接头两侧均可见较明显外露螺纹，按通用口径有预警价值。",
            "上方接头可见但部分边界受前景钢筋遮挡。",
            "结论：可作为“异常预警样本”进入第一版交付。",
        ],
        "boxes": [(0.15, 0.05, 0.84, 0.24), (0.12, 0.39, 0.88, 0.64)],
    },
}


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def manifest_rows() -> dict[str, dict]:
    raw = load_json(MANIFEST_PATH)
    rows = raw["value"] if isinstance(raw, dict) and "value" in raw else raw
    return {row["media_id"]: row for row in rows}


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


FONT_TITLE = load_font(42, bold=True)
FONT_SUBTITLE = load_font(30, bold=True)
FONT_BODY = load_font(25)
FONT_SMALL = load_font(21)
FONT_LABEL = load_font(26, bold=True)


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    bbox = draw.textbbox((0, 0), text, font=font)
    return bbox[2] - bbox[0], bbox[3] - bbox[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for char in text:
        trial = current + char
        if text_size(draw, trial, font)[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = char
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: ImageFont.ImageFont, fill: tuple[int, int, int], max_width: int, line_gap: int = 8) -> int:
    x, y = xy
    for line in wrap_text(draw, text, font, max_width):
        draw.text((x, y), line, font=font, fill=fill)
        y += text_size(draw, line, font)[1] + line_gap
    return y


def norm_to_bbox(norm: tuple[float, float, float, float], width: int, height: int) -> tuple[int, int, int, int]:
    x1, y1, x2, y2 = norm
    return (
        max(0, int(x1 * width)),
        max(0, int(y1 * height)),
        min(width, int(x2 * width)),
        min(height, int(y2 * height)),
    )


def expand_bbox(bbox: tuple[int, int, int, int], width: int, height: int, margin_ratio: float = 0.18) -> tuple[int, int, int, int]:
    x1, y1, x2, y2 = bbox
    bw = x2 - x1
    bh = y2 - y1
    mx = int(bw * margin_ratio)
    my = int(bh * margin_ratio)
    return max(0, x1 - mx), max(0, y1 - my), min(width, x2 + mx), min(height, y2 + my)


def resize_fit(im: Image.Image, max_size: tuple[int, int]) -> Image.Image:
    out = im.copy()
    out.thumbnail(max_size)
    return out


def draw_box(draw: ImageDraw.ImageDraw, bbox: tuple[int, int, int, int], label: str, color: tuple[int, int, int], scale: float = 1.0) -> None:
    x1, y1, x2, y2 = bbox
    width = max(7, int(8 * scale))
    draw.rectangle(bbox, outline=color, width=width)
    label_w, label_h = text_size(draw, label, FONT_LABEL)
    pad = 8
    y_label = max(0, y1 - label_h - pad * 2)
    draw.rectangle((x1, y_label, x1 + label_w + pad * 2, y_label + label_h + pad * 2), fill=color)
    draw.text((x1 + pad, y_label + pad), label, font=FONT_LABEL, fill=COLORS["white"])


def make_board(
    media_id: str,
    source_path: Path,
    scenario: str,
    verdict: str,
    title: str,
    bullets: list[str],
    boxes: list[tuple[int, int, int, int]],
    box_labels: list[str],
    level: str,
) -> Path:
    original = Image.open(source_path).convert("RGB")
    ow, oh = original.size
    color = COLORS[level]

    display = resize_fit(original, (980, 1080))
    scale_x = display.size[0] / ow
    scale_y = display.size[1] / oh
    draw_display = ImageDraw.Draw(display)
    for idx, bbox in enumerate(boxes):
        scaled = tuple(int(v) for v in (bbox[0] * scale_x, bbox[1] * scale_y, bbox[2] * scale_x, bbox[3] * scale_y))
        draw_box(draw_display, scaled, box_labels[idx], color)

    crop_images: list[tuple[str, Image.Image]] = []
    for idx, bbox in enumerate(boxes[:3]):
        crop_box = expand_bbox(bbox, ow, oh, 0.25)
        crop = original.crop(crop_box)
        crop = resize_fit(crop, (650, 280))
        crop_draw = ImageDraw.Draw(crop)
        crop_draw.rectangle((0, 0, crop.size[0] - 1, crop.size[1] - 1), outline=color, width=8)
        crop_images.append((box_labels[idx], crop))

    right_h = 160
    scratch = Image.new("RGB", (700, 100), COLORS["white"])
    scratch_draw = ImageDraw.Draw(scratch)
    for bullet in bullets:
        right_h += 38 * max(1, len(wrap_text(scratch_draw, "• " + bullet, FONT_BODY, 650))) + 22
    right_h += sum(crop.size[1] + 58 for _, crop in crop_images)
    board_h = max(290 + display.size[1] + 60, right_h + 260)
    board_w = 1760
    board = Image.new("RGB", (board_w, board_h), COLORS["light_bg"])
    draw = ImageDraw.Draw(board)

    draw.rectangle((0, 0, board_w, 210), fill=COLORS["dark"])
    draw.text((46, 34), f"{media_id}｜{scenario}", font=FONT_TITLE, fill=COLORS["white"])
    draw.rectangle((46, 112, 46 + 580, 176), fill=color)
    draw.text((68, 126), verdict, font=FONT_SUBTITLE, fill=COLORS["white"])
    draw.text((660, 126), title, font=FONT_SUBTITLE, fill=COLORS["white"])

    left_x, top_y = 46, 250
    board.paste(display, (left_x, top_y))
    draw.rectangle((left_x, top_y, left_x + display.size[0], top_y + display.size[1]), outline=(210, 216, 224), width=2)
    draw.text((left_x, top_y + display.size[1] + 14), "原图加粗标注", font=FONT_SMALL, fill=(85, 95, 108))

    right_x = 1080
    y = 250
    draw.text((right_x, y), "复核结论", font=FONT_SUBTITLE, fill=COLORS["dark"])
    y += 48
    for bullet in bullets:
        y = draw_wrapped(draw, (right_x, y), "• " + bullet, FONT_BODY, COLORS["dark"], 630, 9) + 12
    if crop_images:
        y += 8
        draw.text((right_x, y), "局部放大", font=FONT_SUBTITLE, fill=COLORS["dark"])
        y += 48
    for label, crop in crop_images:
        draw.text((right_x, y), label, font=FONT_LABEL, fill=color)
        y += 34
        board.paste(crop, (right_x, y))
        y += crop.size[1] + 24

    out = CLEAR_ASSETS / f"{media_id}_clear_review.jpg"
    board.save(out, format="JPEG", quality=90, optimize=True)
    return out


def material_review(row: dict, media: dict) -> tuple[str, str, str, list[str], list[tuple[int, int, int, int]], list[str], str]:
    scenario = "钢筋材料点数"
    if row["analysis_status"] == "recapture_required":
        return (
            scenario,
            "不能点数，需补拍",
            "侧面/遮挡视角不能作为最终数量",
            [
                "材料点数只适合端面清楚、无遮挡、能看到每根端头的照片。",
                "这张只能说明现场有一捆钢筋，不能输出可靠数量。",
                "交付时应把这类图直接判为“拍照不合格”。",
            ],
            [],
            [],
            "yellow",
        )
    count = int(row.get("detected_count") or 0)
    return (
        scenario,
        f"候选点数：{count}，需复核",
        "0026可用，旧版5个已撤销",
        [
            "这张是当前唯一适合做材料点数的照片。",
            f"已改为端面候选检测，当前圈出{count}个可见端面候选；旧版只识别5个是漏检，已撤销。",
            "该结果仍需人工点核/补充标注后才能作为验收数量，不建议直接用于收料结算。",
        ],
        [],
        [],
        "yellow",
    )


def concrete_location(bbox: tuple[int, int, int, int], width: int, height: int) -> str:
    x1, y1, x2, y2 = bbox
    cx = (x1 + x2) / 2 / max(width, 1)
    cy = (y1 + y2) / 2 / max(height, 1)
    horizontal = "左" if cx < 0.33 else "中" if cx < 0.66 else "右"
    vertical = "上" if cy < 0.33 else "中" if cy < 0.66 else "下"
    return f"{vertical}{horizontal}"


def concrete_review(row: dict, media: dict) -> tuple[str, str, str, list[str], list[tuple[int, int, int, int]], list[str], str]:
    scenario = "混凝土面表观初筛"
    anomalies = row.get("anomalies", [])[:3]
    boxes = [tuple(int(v) for v in a["bbox"]) for a in anomalies]
    width = int(media.get("width") or 1)
    height = int(media.get("height") or 1)
    labels = []
    type_lines = []
    for idx, anomaly in enumerate(anomalies, 1):
        bbox = tuple(int(v) for v in anomaly["bbox"])
        cn = CLASS_CN.get(anomaly["anomaly_class"], anomaly["anomaly_class"])
        loc = concrete_location(bbox, width, height)
        labels.append(f"疑似{idx}")
        type_lines.append(f"{loc}区域：{cn}")
    bullets = [
        f"本图筛出{len(row.get('anomalies', []))}处疑似表观异常；这里只标最明显的前3处。",
        "当前只能做“疑似位置提醒”，不能替代现场尺量和质量员定性。",
        "如果要验收，需补充构件编号、近景、比例尺/靠尺、处理记录。",
    ]
    bullets.extend(type_lines)
    return (
        scenario,
        "疑似异常，待复核",
        "标位置，不直接判正式缺陷",
        bullets,
        boxes,
        labels,
        "yellow",
    )


def build_records() -> list[dict]:
    CLEAR_ASSETS.mkdir(parents=True, exist_ok=True)
    manifest = manifest_rows()
    records: list[dict] = []

    for row in load_json(MATERIAL_REPORT):
        media = manifest[row["media_id"]]
        scenario, verdict, title, bullets, boxes, labels, level = material_review(row, media)
        source_path = Path(row.get("annotated_image_path") or media["file_path"]) if row["analysis_status"] != "recapture_required" else Path(media["file_path"])
        path = make_board(row["media_id"], source_path, scenario, verdict, title, bullets, boxes, labels, level)
        records.append({"media_id": row["media_id"], "scenario": scenario, "verdict": verdict, "title": title, "bullets": bullets, "board": path, "level": level})

    for row in load_json(COUPLER_REPORT):
        media = manifest[row["media_id"]]
        review = COUPLER_REVIEWS[row["media_id"]]
        source = Path(media["file_path"])
        with Image.open(source) as im:
            w, h = im.size
        boxes = [norm_to_bbox(b, w, h) for b in review["boxes"]]
        labels = [f"观察点{i}" for i in range(1, len(boxes) + 1)]
        path = make_board(
            row["media_id"],
            source,
            "钢筋套筒连接状态",
            review["verdict"],
            review["title"],
            review["bullets"],
            boxes,
            labels,
            review["level"],
        )
        records.append(
            {
                "media_id": row["media_id"],
                "scenario": "钢筋套筒连接状态",
                "verdict": review["verdict"],
                "title": review["title"],
                "bullets": review["bullets"],
                "board": path,
                "level": review["level"],
            }
        )

    concrete_json = load_json(CONCRETE_REPORT)
    for row in concrete_json["results"]:
        media = manifest[row["media_id"]]
        scenario, verdict, title, bullets, boxes, labels, level = concrete_review(row, media)
        path = make_board(row["media_id"], Path(media["file_path"]), scenario, verdict, title, bullets, boxes, labels, level)
        records.append({"media_id": row["media_id"], "scenario": scenario, "verdict": verdict, "title": title, "bullets": bullets, "board": path, "level": level})

    return records


def rel(path: Path) -> str:
    return path.relative_to(DELIVERY_DIR).as_posix()


def build_html(records: list[dict]) -> None:
    sections = {
        "钢筋套筒连接状态": [],
        "钢筋材料点数": [],
        "混凝土面表观初筛": [],
    }
    for record in records:
        sections.setdefault(record["scenario"], []).append(record)

    def card(record: dict) -> str:
        bullet_html = "".join(f"<li>{html.escape(b)}</li>" for b in record["bullets"])
        return f"""
        <article class="card {html.escape(record['level'])}">
          <h3>{html.escape(record['media_id'])}｜{html.escape(record['verdict'])}</h3>
          <p class="sub">{html.escape(record['title'])}</p>
          <ul>{bullet_html}</ul>
          <img src="{html.escape(rel(record['board']))}" alt="{html.escape(record['media_id'])}清晰复核板">
        </article>
        """

    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>AI质监现场照片清晰复核版</title>
  <style>
    body {{ margin:0; font-family:"Microsoft YaHei","PingFang SC",Arial,sans-serif; background:#f4f6f8; color:#16202a; }}
    header {{ background:#17212b; color:white; padding:34px 44px; }}
    header h1 {{ margin:0 0 10px; font-size:30px; }}
    header p {{ max-width:1120px; line-height:1.75; margin:6px 0; }}
    main {{ max-width:1220px; margin:0 auto; padding:26px 22px 60px; }}
    .lead {{ background:white; border:1px solid #d8dee6; border-radius:8px; padding:18px 22px; line-height:1.8; }}
    h2 {{ margin:28px 0 14px; font-size:23px; color:#173b5c; }}
    .card {{ background:white; border:1px solid #d9e0e8; border-radius:8px; padding:18px; margin:16px 0 26px; box-shadow:0 1px 2px rgba(0,0,0,.04); }}
    .card h3 {{ margin:0 0 6px; font-size:20px; }}
    .card.red h3 {{ color:#b92d24; }}
    .card.yellow h3 {{ color:#a85f05; }}
    .card.green h3 {{ color:#22704d; }}
    .sub {{ margin:0 0 8px; font-weight:700; }}
    ul {{ line-height:1.75; margin:8px 0 14px 22px; }}
    img {{ width:100%; height:auto; display:block; border-radius:6px; border:1px solid #e2e7ef; }}
    .rule {{ font-size:14px; color:#55616f; }}
  </style>
</head>
<body>
<header>
  <h1>AI质监现场照片清晰复核版</h1>
  <p>这版按“人能看懂”重做：套筒不做数量清点，只判断连接状态、露丝风险、是否需要补拍；每张图都给大结论、粗框和局部放大。</p>
  <p>生成时间：2026-05-25。当前结论用于会议评估和样板验收，不替代正式质量验收记录。</p>
</header>
<main>
  <section class="lead">
    <h2>先说结论</h2>
    <p><b>套筒：</b>这组图不应做“数量清点”。5张里，0019像单端/未完整连接，0020和0022太远只能定位，0021和0023近景较清楚且有疑似露丝偏多风险，适合做第一版异常预警样本。</p>
    <p><b>钢筋点数：</b>只有端面图0026适合做这个场景；旧算法只输出5个点位已撤销，新版圈出端面候选点，但仍需人工点核/补充标注后才能作为验收数量。其他侧面/遮挡图直接判为拍照不合格。</p>
    <p><b>混凝土面：</b>当前只能做表观疑似位置提醒，不能直接判蜂窝、麻面、裂缝等正式缺陷；需要近景、比例尺/靠尺和构件编号配合。</p>
    <p class="rule">默认规范口径：套筒露丝按JGJ 107-2016相关公开摘录的“单侧外露螺纹不宜超过2p”作为预警线；混凝土外观类别参考GB 50204-2015现浇结构外观质量缺陷分类。</p>
  </section>

  <section>
    <h2>一、钢筋套筒连接状态，不做数量清点</h2>
    {''.join(card(r) for r in sections['钢筋套筒连接状态'])}
  </section>

  <section>
    <h2>二、钢筋材料点数，只接受端面照片</h2>
    {''.join(card(r) for r in sections['钢筋材料点数'])}
  </section>

  <section>
    <h2>三、混凝土面表观初筛，只标疑似位置</h2>
    {''.join(card(r) for r in sections['混凝土面表观初筛'])}
  </section>
</main>
</body>
</html>
"""
    HTML_OUT.write_text(html_text, encoding="utf-8")


def build_docx(records: list[dict]) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    for style_name, size in [("Normal", 10.5), ("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI质监现场照片清晰复核版").bold = True
    subtitle = doc.add_paragraph("生成时间：2026-05-25｜套筒不做数量清点｜粗框+局部放大版")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("先说结论", level=1)
    summary = [
        "套筒：这组图不应做数量清点。0019像单端/未完整连接，0020和0022太远只能定位，0021和0023疑似露丝偏多，适合做异常预警样本。",
        "钢筋点数：只有0026端面图适合做这个场景；旧算法只输出5个点位已撤销，新版圈出端面候选点，但仍需人工点核/补充标注后才能作为验收数量。其他侧面/遮挡图判为拍照不合格。",
        "混凝土面：当前只能做表观疑似位置提醒，不能直接判正式缺陷；需要近景、比例尺/靠尺和构件编号配合。",
    ]
    for item in summary:
        doc.add_paragraph(item, style="List Bullet")

    overview = doc.add_table(rows=1, cols=4)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = overview.rows[0].cells
    for idx, text in enumerate(["图片", "场景", "结论", "说明"]):
        header_cells[idx].text = text
    for record in records:
        row = overview.add_row().cells
        row[0].text = record["media_id"]
        row[1].text = record["scenario"]
        row[2].text = record["verdict"]
        row[3].text = record["title"]

    for scenario in ["钢筋套筒连接状态", "钢筋材料点数", "混凝土面表观初筛"]:
        doc.add_heading(scenario, level=1)
        for record in [r for r in records if r["scenario"] == scenario]:
            doc.add_heading(f"{record['media_id']}｜{record['verdict']}", level=2)
            doc.add_paragraph(record["title"])
            for bullet in record["bullets"]:
                doc.add_paragraph(bullet, style="List Bullet")
            doc.add_picture(str(record["board"]), width=Inches(7.1))

    doc.save(DOCX_OUT)


def main() -> None:
    records = build_records()
    build_html(records)
    build_docx(records)
    print(f"HTML: {HTML_OUT}")
    print(f"DOCX: {DOCX_OUT}")
    print(f"clear boards: {len(list(CLEAR_ASSETS.glob('*_clear_review.jpg')))}")


if __name__ == "__main__":
    main()
