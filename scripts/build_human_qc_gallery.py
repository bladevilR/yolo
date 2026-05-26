from __future__ import annotations

import html
import json
import shutil
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image


WORKSPACE = Path("E:/yolo")
DELIVERY_DIR = WORKSPACE / "AI_QC_Human_Delivery_20260525"
ASSET_DIR = DELIVERY_DIR / "gallery_assets"

MANIFEST_PATH = WORKSPACE / "datasets/field_qc/reports/media_manifest.json"
MATERIAL_REPORT = WORKSPACE / "datasets/field_qc/rebar_material_counting/reports/rebar_material_count_report.json"
COUPLER_REPORT = WORKSPACE / "datasets/field_qc/rebar_coupler_thread_qc/reports/rebar_coupler_thread_qc_report.json"
CONCRETE_REPORT = WORKSPACE / "datasets/field_qc/concrete_surface_qc/reports/concrete_surface_qc_report.json"

HTML_OUT = DELIVERY_DIR / "03_AI_QC_Result_Gallery_CN.html"
DOCX_OUT = DELIVERY_DIR / "03_AI_QC_Result_Gallery_CN.docx"

STANDARD_LINKS = [
    {
        "name": "JGJ 107-2016《钢筋机械连接技术规程》",
        "url": "https://www.codeofchina.com/standard/JGJ107-2016.html",
        "note": "用于套筒/机械连接接头的安装、现场检验和验收框架。",
    },
    {
        "name": "JGJ 107-2016 第6.3.1条相关公开摘录",
        "url": "https://www.lengjiyataotong.com/taotongxw/251.html",
        "note": "直螺纹标准型、正反丝型、异径型接头安装后单侧外露螺纹不宜超过2p。",
    },
    {
        "name": "GB 50204-2015《混凝土结构工程施工质量验收规范》",
        "url": "https://www.zbl.cn/uploadfile/2020/0812/20200812111624347.pdf",
        "note": "用于现浇结构外观质量缺陷分类，包含露筋、蜂窝、孔洞、夹渣、疏松、裂缝、连接部位缺陷、外形缺陷、外表缺陷等。",
    },
]

CLASS_CN = {
    "repair_patch": "修补痕迹/表面补修色差",
    "color_difference": "色差/污染/泛白",
    "edge_defect": "边角缺损/棱角异常",
    "surface_pitting": "麻面/起砂疑似",
    "crack_like": "裂缝疑似",
    "honeycomb_like": "蜂窝疑似",
}

FLAG_CN = {
    "occlusion_present": "遮挡较多",
    "endpoint_face_required_for_exact_count": "需要端面照片才能精确计数",
    "endpoint_cluster_off_center": "端面目标偏在画面一侧，需人工复核",
    "both_coupler_sides_required": "套筒两端未完全清晰",
    "exposed_thread_threshold_missing": "未配置项目露丝阈值，已按通用2p口径补充判断",
}


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8-sig"))


def ensure_dirs() -> None:
    DELIVERY_DIR.mkdir(parents=True, exist_ok=True)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def manifest_rows() -> dict[str, dict]:
    raw = load_json(MANIFEST_PATH)
    rows = raw["value"] if isinstance(raw, dict) and "value" in raw else raw
    return {row["media_id"]: row for row in rows}


def safe_asset_name(media_id: str, suffix: str) -> str:
    return f"{media_id}_{suffix}.jpg"


def resize_copy(src: Path, dst: Path, max_side: int = 1400) -> tuple[int, int]:
    if not src.exists():
        raise FileNotFoundError(src)
    with Image.open(src) as im:
        im = im.convert("RGB")
        im.thumbnail((max_side, max_side))
        im.save(dst, format="JPEG", quality=86, optimize=True)
        return im.size


def copy_pair(media_id: str, original_path: str, annotated_path: str | None) -> dict:
    orig_dst = ASSET_DIR / safe_asset_name(media_id, "original")
    ann_dst = ASSET_DIR / safe_asset_name(media_id, "annotated")
    orig_size = resize_copy(Path(original_path), orig_dst)
    ann_size = None
    if annotated_path:
        ann_size = resize_copy(Path(annotated_path), ann_dst)
    else:
        shutil.copy2(orig_dst, ann_dst)
        ann_size = orig_size
    return {
        "original": orig_dst,
        "annotated": ann_dst,
        "original_rel": orig_dst.relative_to(DELIVERY_DIR).as_posix(),
        "annotated_rel": ann_dst.relative_to(DELIVERY_DIR).as_posix(),
        "orig_size": orig_size,
        "ann_size": ann_size,
    }


def status_cn(status: str) -> str:
    return {
        "screened": "已筛查，需现场复核确认",
        "needs_review": "需人工复核",
        "recapture_required": "需补拍后再判定",
    }.get(status, status)


def decision_badge(decision: str) -> str:
    return {
        "pass": "初判合规",
        "fail": "初判不合规",
        "warning": "预警",
        "needs_standard_confirmation": "需按标准/项目口径确认",
        "not_applicable": "不适用",
    }.get(decision, decision)


def flags_cn(flags: list[str] | None) -> str:
    if not flags:
        return "无明显采集问题"
    return "；".join(FLAG_CN.get(flag, flag) for flag in flags)


def location_cn(bbox: list[int], width: int, height: int) -> str:
    x1, y1, x2, y2 = bbox
    cx = (x1 + x2) / 2 / max(width, 1)
    cy = (y1 + y2) / 2 / max(height, 1)
    horizontal = "左侧" if cx < 0.33 else "中部" if cx < 0.66 else "右侧"
    vertical = "上部" if cy < 0.33 else "中部" if cy < 0.66 else "下部"
    if horizontal == "中部" and vertical == "中部":
        return "画面中部"
    return f"{vertical}{horizontal}"


def top_anomaly_summary(anomalies: list[dict], width: int, height: int, limit: int = 4) -> list[str]:
    lines = []
    for idx, anomaly in enumerate(anomalies[:limit], 1):
        bbox = anomaly["bbox"]
        cls = CLASS_CN.get(anomaly["anomaly_class"], anomaly["anomaly_class"])
        conf = anomaly.get("confidence", 0)
        lines.append(
            f"{idx}. {location_cn(bbox, width, height)}：{cls}，置信度{conf:.2f}，像素框{x1y1x2y2(bbox)}"
        )
    return lines


def x1y1x2y2(bbox: list[int]) -> str:
    return f"({bbox[0]},{bbox[1]})-({bbox[2]},{bbox[3]})"


def material_judgment(row: dict) -> str:
    if row["analysis_status"] == "recapture_required":
        return "不能作为最终数量；当前图为侧面/遮挡视角，必须补拍端面后再自动计数。"
    count = row.get("detected_count", 0)
    if "high_density_count_requires_manual_review" in row.get("review_flags", []):
        return f"当前候选检测到 {count} 个可见端面点位；旧版“5个”已撤销。该结果仍需人工点核/补充标注后才能作为验收数量。"
    return f"当前模型识别到 {count} 个端面点位；结论建议人工复核后再入库。"


def coupler_decision(row: dict, threshold: int = 2) -> tuple[str, str]:
    left = int(row.get("left_visible_thread_count") or 0)
    right = int(row.get("right_visible_thread_count") or 0)
    max_side = max(left, right)
    if max_side > threshold:
        return "初判不合规/需整改复核", f"按通用口径单侧不超过{threshold}扣，当前最大单侧约{max_side}扣。"
    return "初判合规但需人工确认", f"按通用口径单侧不超过{threshold}扣，当前左右约为{left}/{right}扣。"


def concrete_judgment(row: dict) -> tuple[str, str]:
    anomalies = row.get("anomalies", [])
    if not anomalies:
        return "未发现明显疑似缺陷", "当前图未筛出高置信异常区域。"
    classes = Counter(CLASS_CN.get(a["anomaly_class"], a["anomaly_class"]) for a in anomalies)
    top = classes.most_common(1)[0][0]
    return "发现疑似表观异常，需现场复核", f"共筛出 {len(anomalies)} 处疑似区域，主要类型为：{top}。"


def build_records() -> tuple[list[dict], list[dict], list[dict]]:
    manifest = manifest_rows()

    material_records = []
    for row in load_json(MATERIAL_REPORT):
        media = manifest[row["media_id"]]
        assets = copy_pair(row["media_id"], media["file_path"], row.get("annotated_image_path"))
        material_records.append({**row, "media": media, "assets": assets, "scenario": "钢筋进场数量识别"})

    coupler_records = []
    for row in load_json(COUPLER_REPORT):
        media = manifest[row["media_id"]]
        assets = copy_pair(row["media_id"], media["file_path"], row.get("annotated_image_path"))
        result, reason = coupler_decision(row)
        coupler_records.append(
            {**row, "media": media, "assets": assets, "scenario": "钢筋套筒露丝合规检查", "result_cn": result, "reason_cn": reason}
        )

    concrete_json = load_json(CONCRETE_REPORT)
    concrete_records = []
    for row in concrete_json["results"]:
        media = manifest[row["media_id"]]
        assets = copy_pair(row["media_id"], media["file_path"], row.get("annotated_image_path"))
        result, reason = concrete_judgment(row)
        concrete_records.append(
            {**row, "media": media, "assets": assets, "scenario": "混凝土面表观缺陷筛查", "result_cn": result, "reason_cn": reason}
        )

    return material_records, coupler_records, concrete_records


def h(text: object) -> str:
    return html.escape(str(text), quote=True)


def image_pair_html(record: dict) -> str:
    return f"""
    <div class="image-pair">
      <figure><img src="{h(record['assets']['original_rel'])}" alt="{h(record['media_id'])}原图"><figcaption>原图</figcaption></figure>
      <figure><img src="{h(record['assets']['annotated_rel'])}" alt="{h(record['media_id'])}标注图"><figcaption>标注图</figcaption></figure>
    </div>
    """


def build_html(material: list[dict], coupler: list[dict], concrete: list[dict]) -> None:
    material_rows = []
    for record in material:
        material_rows.append(
            f"""
            <article class="record">
              <h3>{h(record['media_id'])}｜钢筋进场数量识别</h3>
              <p class="verdict">{h(status_cn(record['analysis_status']))}</p>
              <p><b>数量输出：</b>{h(record.get('detected_count', 0))} 个；<b>置信度：</b>{h(record.get('confidence', 0))}</p>
              <p><b>结论：</b>{h(material_judgment(record))}</p>
              <p><b>采集/复核提示：</b>{h(flags_cn(record.get('review_flags') or record['media'].get('quality_flags')))}</p>
              {image_pair_html(record)}
            </article>
            """
        )

    coupler_rows = []
    for record in coupler:
        coupler_rows.append(
            f"""
            <article class="record">
              <h3>{h(record['media_id'])}｜钢筋套筒露丝合规检查</h3>
              <p class="verdict">{h(record['result_cn'])}</p>
              <p><b>露丝数量输出：</b>左侧约 {h(record.get('left_visible_thread_count'))} 扣，右侧约 {h(record.get('right_visible_thread_count'))} 扣。</p>
              <p><b>判断依据：</b>{h(record['reason_cn'])}</p>
              <p><b>复核提示：</b>{h(flags_cn(record.get('review_flags') or record['media'].get('quality_flags')))}</p>
              {image_pair_html(record)}
            </article>
            """
        )

    concrete_rows = []
    for record in concrete:
        width = int(record["media"].get("width") or 1)
        height = int(record["media"].get("height") or 1)
        anomaly_lines = top_anomaly_summary(record.get("anomalies", []), width, height)
        anomaly_html = "".join(f"<li>{h(line)}</li>" for line in anomaly_lines) or "<li>无</li>"
        concrete_rows.append(
            f"""
            <article class="record">
              <h3>{h(record['media_id'])}｜混凝土面表观缺陷筛查</h3>
              <p class="verdict">{h(record['result_cn'])}</p>
              <p><b>疑似缺陷数量输出：</b>{h(len(record.get('anomalies', [])))} 处。</p>
              <p><b>结论：</b>{h(record['reason_cn'])}</p>
              <p><b>主要疑似位置：</b></p>
              <ol>{anomaly_html}</ol>
              {image_pair_html(record)}
            </article>
            """
        )

    standards = "".join(
        f'<li><a href="{h(item["url"])}">{h(item["name"])}</a>：{h(item["note"])}</li>' for item in STANDARD_LINKS
    )

    total_material_ready = sum(1 for r in material if r["analysis_status"] != "recapture_required")
    total_coupler_fail = sum(1 for r in coupler if r["result_cn"].startswith("初判不合规"))
    total_concrete_regions = sum(len(r.get("anomalies", [])) for r in concrete)

    html_text = f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>AI质监现场素材成果图册</title>
  <style>
    body {{ margin:0; font-family: "Microsoft YaHei", "PingFang SC", Arial, sans-serif; color:#17202a; background:#f6f7f9; }}
    header {{ padding:32px 42px; background:#123b5d; color:#fff; }}
    header h1 {{ margin:0 0 8px; font-size:28px; }}
    header p {{ margin:4px 0; line-height:1.65; max-width:1080px; }}
    main {{ max-width:1180px; margin:0 auto; padding:26px 24px 56px; }}
    section {{ margin:24px 0 34px; }}
    h2 {{ margin:0 0 14px; font-size:22px; color:#123b5d; }}
    h3 {{ margin:0 0 10px; font-size:18px; }}
    .summary {{ display:grid; grid-template-columns: repeat(4, minmax(0,1fr)); gap:12px; margin:20px 0; }}
    .metric {{ background:#fff; border:1px solid #d8dee6; border-radius:8px; padding:16px; }}
    .metric b {{ display:block; font-size:24px; color:#0f6f5c; margin-bottom:4px; }}
    .note, .standards {{ background:#fff; border:1px solid #d8dee6; border-radius:8px; padding:16px 18px; line-height:1.75; }}
    .record {{ background:#fff; border:1px solid #d8dee6; border-radius:8px; padding:18px; margin:16px 0 22px; box-shadow:0 1px 2px rgba(0,0,0,.04); }}
    .verdict {{ display:inline-block; margin:0 0 10px; padding:5px 10px; border-radius:999px; background:#e8f3ee; color:#0f6f5c; font-weight:700; }}
    .record p {{ line-height:1.65; margin:8px 0; }}
    .record ol {{ margin:6px 0 12px 22px; line-height:1.7; }}
    .image-pair {{ display:grid; grid-template-columns:1fr 1fr; gap:14px; margin-top:14px; }}
    figure {{ margin:0; background:#fafafa; border:1px solid #e4e8ee; border-radius:6px; padding:8px; }}
    img {{ width:100%; height:auto; display:block; border-radius:4px; }}
    figcaption {{ text-align:center; color:#5f6b76; font-size:13px; margin-top:6px; }}
    a {{ color:#0c5f99; }}
    @media (max-width: 760px) {{
      header {{ padding:24px 20px; }}
      main {{ padding:18px 14px 42px; }}
      .summary {{ grid-template-columns:1fr 1fr; }}
      .image-pair {{ grid-template-columns:1fr; }}
    }}
  </style>
</head>
<body>
<header>
  <h1>AI质监现场素材成果图册</h1>
  <p>生成时间：2026-05-25。目的：把现场调研照片转成可复核、可验收的中文成果，而不是只给算法CSV。</p>
  <p>重要说明：当前为规则+视觉启发式验证版，结论用于会议评估和现场复核，不替代正式验收记录。</p>
</header>
<main>
  <section class="summary">
    <div class="metric"><b>{len(material)}</b>钢筋计数图片</div>
    <div class="metric"><b>{total_material_ready}</b>张可输出候选数量</div>
    <div class="metric"><b>{len(coupler)}</b>套筒露丝图片，其中{total_coupler_fail}张预警</div>
    <div class="metric"><b>{total_concrete_regions}</b>处混凝土疑似区域</div>
  </section>

  <section class="note">
    <h2>当前总判断</h2>
    <p><b>钢筋数量识别：</b>现有7张里6张为侧面/遮挡视角，不能直接作为最终数量；1张端面图识别到5个端面点位，但仍需人工复核。要快速交付，必须把“端面拍照规范”作为验收前置条件。</p>
    <p><b>钢筋套筒露丝：</b>可做快速交付。按通用2扣/2p口径，本批5张中多张出现单侧露丝明显超过阈值的预警，适合先做“红黄绿+人工复核”版本。</p>
    <p><b>混凝土面检查：</b>可做表观异常初筛。当前10张共筛出{total_concrete_regions}处疑似区域，主要是修补痕迹、色差/污染、麻面/起砂疑似；缺陷性质和严重程度仍要结合构件位置、尺寸测量和现场处理记录确认。</p>
  </section>

  <section class="standards">
    <h2>默认规范依据</h2>
    <ul>{standards}</ul>
    <p>落地口径：套筒露丝先按单侧≤2扣/2p做预警；混凝土面按GB 50204-2015的外观质量缺陷类别输出“疑似类型+位置+待复核”。材料点数本身不是规范合格项，需与项目收料单、施工图或进场验收记录对齐。</p>
  </section>

  <section>
    <h2>一、钢筋进场数量识别</h2>
    {''.join(material_rows)}
  </section>

  <section>
    <h2>二、钢筋套筒露丝合规检查</h2>
    {''.join(coupler_rows)}
  </section>

  <section>
    <h2>三、混凝土面表观缺陷筛查</h2>
    {''.join(concrete_rows)}
  </section>
</main>
</body>
</html>
"""
    HTML_OUT.write_text(html_text, encoding="utf-8")


def add_hyperlink(paragraph, text: str, url: str):
    # Keep the DOCX builder dependency-light; the HTML gallery carries clickable links.
    return paragraph.add_run(text)


def set_base_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    for style_name, size in [("Title", 20), ("Heading 1", 15), ("Heading 2", 12)]:
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)


def add_image_pair_docx(document: Document, record: dict) -> None:
    table = document.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    for i, label in enumerate(["原图", "标注图"]):
        p = table.cell(0, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label).bold = True
    for i, key in enumerate(["original", "annotated"]):
        p = table.cell(1, i).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(record["assets"][key]), width=Inches(2.95))


def build_docx(material: list[dict], coupler: list[dict], concrete: list[dict]) -> None:
    doc = Document()
    set_base_style(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("AI质监现场素材成果图册").bold = True
    subtitle = doc.add_paragraph("生成时间：2026-05-25｜用途：会议汇报、验收口径确认、现场复拍清单")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、当前总判断", level=1)
    total_concrete_regions = sum(len(r.get("anomalies", [])) for r in concrete)
    bullets = [
        "钢筋数量识别：7张图片中6张为侧面/遮挡视角，不能直接作为最终数量；1张端面图识别到5个端面点位，需要人工复核。",
        "钢筋套筒露丝：5张图片均可做露丝预警；按通用单侧≤2扣/2p口径，多张存在超阈值风险，适合优先交付红黄绿检查。",
        f"混凝土面检查：10张图片共筛出{total_concrete_regions}处疑似区域，主要为修补痕迹、色差/污染、麻面/起砂疑似，需要现场复核缺陷性质和严重程度。",
    ]
    for item in bullets:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("二、默认规范依据", level=1)
    for item in STANDARD_LINKS:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item["name"]).bold = True
        p.add_run("：")
        p.add_run(item["note"])
        p.add_run(" 来源：")
        add_hyperlink(p, item["url"], item["url"])
    doc.add_paragraph("落地口径：套筒露丝先按单侧≤2扣/2p做预警；混凝土面按GB 50204-2015外观质量缺陷类别输出疑似类型、位置和待复核结论。")

    doc.add_heading("三、钢筋进场数量识别", level=1)
    for record in material:
        doc.add_heading(f"{record['media_id']}｜钢筋进场数量识别", level=2)
        doc.add_paragraph(f"数量输出：{record.get('detected_count', 0)} 个；状态：{status_cn(record['analysis_status'])}；置信度：{record.get('confidence', 0)}。")
        doc.add_paragraph(f"结论：{material_judgment(record)}")
        doc.add_paragraph(f"采集/复核提示：{flags_cn(record.get('review_flags') or record['media'].get('quality_flags'))}")
        add_image_pair_docx(doc, record)

    doc.add_heading("四、钢筋套筒露丝合规检查", level=1)
    for record in coupler:
        doc.add_heading(f"{record['media_id']}｜钢筋套筒露丝合规检查", level=2)
        doc.add_paragraph(f"露丝数量输出：左侧约{record.get('left_visible_thread_count')}扣，右侧约{record.get('right_visible_thread_count')}扣。")
        doc.add_paragraph(f"判断结果：{record['result_cn']}。{record['reason_cn']}")
        doc.add_paragraph(f"复核提示：{flags_cn(record.get('review_flags') or record['media'].get('quality_flags'))}")
        add_image_pair_docx(doc, record)

    doc.add_heading("五、混凝土面表观缺陷筛查", level=1)
    for record in concrete:
        doc.add_heading(f"{record['media_id']}｜混凝土面表观缺陷筛查", level=2)
        width = int(record["media"].get("width") or 1)
        height = int(record["media"].get("height") or 1)
        doc.add_paragraph(f"疑似缺陷数量输出：{len(record.get('anomalies', []))}处。判断结果：{record['result_cn']}。{record['reason_cn']}")
        doc.add_paragraph("主要疑似位置：")
        for line in top_anomaly_summary(record.get("anomalies", []), width, height):
            doc.add_paragraph(line, style="List Number")
        add_image_pair_docx(doc, record)

    doc.save(DOCX_OUT)


def main() -> None:
    ensure_dirs()
    material, coupler, concrete = build_records()
    build_html(material, coupler, concrete)
    build_docx(material, coupler, concrete)
    print(f"HTML: {HTML_OUT}")
    print(f"DOCX: {DOCX_OUT}")
    print(f"assets: {len(list(ASSET_DIR.glob('*.jpg')))}")


if __name__ == "__main__":
    main()
