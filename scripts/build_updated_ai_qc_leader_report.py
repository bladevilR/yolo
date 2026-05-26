from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("E:/yolo")
DELIVERY = ROOT / "AI_QC_Human_Delivery_20260525"
OUT_MAIN = DELIVERY / "01_AI_QC_Demo_Leader_Brief_CN.docx"
OUT_VERSIONED = DELIVERY / "06_AI_QC_Leader_Report_Updated_CN.docx"

IMG_COUNT = DELIVERY / "clear_review_assets/field-qc-0026_clear_review.jpg"
IMG_COUPLER = DELIVERY / "clear_review_assets/field-qc-0023_clear_review.jpg"
IMG_CONCRETE = DELIVERY / "clear_review_assets/field-qc-0005_clear_review.jpg"

INK = RGBColor(24, 32, 43)
BLUE = RGBColor(31, 88, 135)
RED = RGBColor(174, 43, 36)
GREEN = RGBColor(30, 110, 76)
MUTED = RGBColor(92, 102, 112)
LIGHT_BLUE = "E8F1FA"
LIGHT_YELLOW = "FFF5D6"
LIGHT_RED = "FCECEC"
LIGHT_GREEN = "E9F5EF"
BORDER = "C8D2E0"


def set_run(run, size=10.5, bold=False, color=INK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def borders(table, color=BORDER, size="4"):
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = tbl_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def para(doc, text="", size=10.5, bold=False, color=INK, before=0, after=6, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    r = p.add_run(text)
    set_run(r, size=15 if level == 1 else 12.5, bold=True, color=BLUE)
    return p


def bullet(doc, text, color=INK):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run(r, size=10, color=color)
    return p


def callout(doc, title, body, fill=LIGHT_YELLOW, title_color=RED):
    table = doc.add_table(rows=1, cols=1)
    table_widths(table, [7.15])
    borders(table, color="D6B656", size="8")
    cell = table.cell(0, 0)
    shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.2
    r1 = p.add_run(title + "：")
    set_run(r1, size=10.5, bold=True, color=title_color)
    r2 = p.add_run(body)
    set_run(r2, size=10.5, color=INK)
    para(doc, "", after=2)


def cell_text(cell, text, bold=False, fill=None, color=INK, size=9):
    cell.text = ""
    if fill:
        shade(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table_widths(table, widths)
    borders(table)
    for idx, header in enumerate(headers):
        cell_text(table.cell(0, idx), header, bold=True, fill=LIGHT_BLUE, size=9.2)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell_text(cells[idx], value, size=8.8)
    table_widths(table, widths)
    borders(table)
    return table


def add_image(doc, path: Path, caption: str):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(6.9))
    cap = para(doc, caption, size=8.8, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    return cap


def set_styles(doc: Document):
    for style_name, size in (("Normal", 10.5), ("Title", 20), ("Heading 1", 15), ("Heading 2", 12.5)):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_doc() -> Document:
    doc = Document()
    set_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    title = para(doc, "AI质监现场需求评估与交付计划（更新版）", size=19, bold=True, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
    subtitle = para(doc, "生成时间：2026-05-25｜用于领导汇报、建设公司会议、下一步训练计划确认", size=9.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)

    callout(
        doc,
        "核心口径",
        "当前可以汇报为“AI质监辅助评估 Demo 已完成一轮现场素材验证”，但不能汇报为“钢筋点数模型已训练完成”。0026 的 118 是端面候选圈，不是验收数量；旧版 5 是漏检，已撤销。",
        fill=LIGHT_RED,
        title_color=RED,
    )

    heading(doc, "一、当前可汇报判断", 1)
    add_table(
        doc,
        ["场景", "当前判断", "能否快速交付", "会议口径"],
        [
            ["钢筋材料点数", "方向可行，但训练闭环缺失；当前只有 0026 可作为端面样本", "暂不作为正式交付场景", "已验证端面照片具备点数条件，但需要补样本和重新训练"],
            ["钢筋套筒连接状态/露丝风险", "可做辅助复核；不做数量清点", "可以优先做 Demo", "输出正常/疑似露丝偏多/看不清需补拍"],
            ["混凝土面表观异常", "可做疑似位置提醒", "可以优先做 Demo", "只标疑似异常位置，不直接判正式缺陷"],
        ],
        [1.35, 2.25, 1.45, 2.15],
    )

    heading(doc, "二、已经完成的东西", 1)
    for item in [
        "现场素材归类：混凝土面、套筒连接、钢筋材料、不可分配素材。",
        "清晰复核版图册：每张图有粗框、局部放大、中文结论。",
        "套筒多模态复核：不再当数量清点，而是判断连接状态、露丝风险和是否需要补拍。",
        "混凝土面疑似区域标注：能给出位置提醒，但不做正式缺陷定性。",
        "钢筋点数问题修正：0026 旧版只识别 5 个是漏检，已撤销。",
        "训练流程重评估：当前没有钢筋端面点数训练集，不能说模型已训练完成。",
    ]:
        bullet(doc, item)

    heading(doc, "三、钢筋材料点数：必须调整口径", 1)
    bullet(doc, "0015-0018、0024、0025 是侧面/遮挡图，不能用于最终点数。")
    bullet(doc, "0026 是唯一可用端面图；当前圈出 118 个可见端面候选。")
    bullet(doc, "118 不是最终验收数量，必须人工点核；旧版 5 是算法漏检，不能再对外使用。", color=RED)
    callout(
        doc,
        "建议说法",
        "钢筋材料点数场景具备可行性，但当前还没形成训练集。0026 已生成 118 个端面候选圈，作为人工预标注起点；下一步要补拍端面样本、人工标注真值、重新训练并按计数误差验收。",
        fill=LIGHT_YELLOW,
        title_color=BLUE,
    )
    add_image(doc, IMG_COUNT, "0026 钢筋端面点数复核板：118 为候选圈，需人工点核后才能作为验收数量。")

    heading(doc, "四、套筒场景：优先 Demo，但不是数量清点", 1)
    add_table(
        doc,
        ["图片", "判断", "说明"],
        [
            ["field-qc-0019", "不能作为完整接头验收", "更像单端/未完整连接状态"],
            ["field-qc-0020", "只能定位，不能验收", "远景多接头，套筒太小，需要逐个近景补拍"],
            ["field-qc-0021", "疑似露丝偏多", "近景可看，但有遮挡，建议复核"],
            ["field-qc-0022", "只能定位，不能验收", "多接头中远景，细节不足"],
            ["field-qc-0023", "疑似露丝偏多", "近景较清楚，两侧露丝可见"],
        ],
        [1.45, 2.0, 3.7],
    )
    bullet(doc, "第一版建议做“连接状态辅助复核”：正常、疑似异常、看不清需补拍。")
    bullet(doc, "露丝判断先按通用“单侧外露螺纹不宜超过 2p/约 2 扣”作为预警线，最终阈值需项目确认。")
    add_image(doc, IMG_COUPLER, "0023 套筒复核板：用于露丝风险预警，不用于数量清点。")

    heading(doc, "五、混凝土面：做疑似位置提醒", 1)
    bullet(doc, "能圈出疑似异常区域，并给出位置、粗框、局部放大和中文说明。")
    bullet(doc, "需要构件编号、近景、比例尺/靠尺和处理记录配合。")
    bullet(doc, "当前不能直接定性为蜂窝、麻面、裂缝、露筋等正式缺陷，也不能替代质量员验收签字。", color=RED)
    add_image(doc, IMG_CONCRETE, "混凝土面复核板：只做疑似位置提醒，不直接判正式缺陷。")

    heading(doc, "六、训练流程重评估", 1)
    add_table(
        doc,
        ["步骤", "要做什么", "输出"],
        [
            ["1", "明确标注规则：可见端面、半遮挡、背景钢筋、ignore 区域", "标注规范"],
            ["2", "补拍端面正样本 80-120 张；不可计数负样本 40-60 张", "训练素材"],
            ["3", "用候选检测做预标注，人工删错补漏", "真值标签"],
            ["4", "训练第一版点数模型，建议高分辨率或切片推理", "模型权重"],
            ["5", "独立测试集验收，不能同一捆相似照片混入训练和测试", "测试报告"],
            ["6", "按 MAE、MAPE、±1/±3/±5 根准确率、拒识率验收", "是否可交付结论"],
        ],
        [0.65, 4.25, 2.25],
    )

    heading(doc, "七、下一步计划", 1)
    add_table(
        doc,
        ["阶段", "工作", "输出物"],
        [
            ["会前", "用 04 清晰复核版汇报现状；不要再用旧 CSV/旧 5 个点数结果", "领导汇报材料、清晰复核图册"],
            ["会上", "确认套筒露丝阈值、混凝土缺陷口径、钢筋端面补拍规范", "会议确认清单"],
            ["会后 1-2 天", "补拍端面照片、套筒近景、混凝土缺陷正样本", "现场补样素材"],
            ["会后 3-5 天", "完成点数真值标注和训练集划分", "标注数据集、训练配置"],
            ["下一版 Demo", "训练点数模型并输出独立测试误差", "模型权重、测试报告、失败案例"],
        ],
        [1.15, 4.15, 1.85],
    )

    heading(doc, "八、领导版结论", 1)
    callout(
        doc,
        "建议结论",
        "当前项目值得继续推进，但应从“正式自动验收”降级为“辅助质检 Demo + 补样训练 + 规则确认”。先把套筒和混凝土两个场景做成可演示、可复核的第一版；钢筋点数作为第二阶段训练任务，等端面样本和标注真值到位后再做验收承诺。",
        fill=LIGHT_GREEN,
        title_color=GREEN,
    )

    return doc


def main():
    doc = build_doc()
    doc.save(OUT_MAIN)
    doc.save(OUT_VERSIONED)
    print(f"updated: {OUT_MAIN}")
    print(f"versioned: {OUT_VERSIONED}")


if __name__ == "__main__":
    main()
